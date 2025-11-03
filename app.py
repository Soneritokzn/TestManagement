from flask import Flask, render_template, request, jsonify, send_from_directory, send_file, redirect, url_for
from models import (
    db, TestCase, Step, TestCaseComment, Attachment, TestCaseTemplate, 
    TemplateStep, TestRun, TestCaseExecution, TestCaseVersion, VersionStep,
    TestStatus, Priority
)
from docx import Document
import os
import json
from datetime import datetime
import pandas as pd
from werkzeug.utils import secure_filename
import uuid

app = Flask(__name__, instance_relative_config=True)

# Ensure instance folder exists
INSTANCE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance')
if not os.path.exists(INSTANCE_PATH):
    os.makedirs(INSTANCE_PATH)

app.config['SQLALCHEMY_DATABASE_URI'] = f'sqlite:///{os.path.join(INSTANCE_PATH, "database.db").replace(os.sep, "/")}'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size

db.init_app(app)

UPLOAD_FOLDER = 'exports'
ATTACHMENT_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf', 'doc', 'docx', 'txt'}

for folder in [UPLOAD_FOLDER, ATTACHMENT_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

with app.app_context():
    db.create_all()

# --- FRONTEND ROUTES ---
@app.route('/')
def index():
    return render_template('index.html')

@app.route('/dashboard')
def dashboard():
    return render_template('index.html')  # Dashboard will be part of the main page

# --- API: DASHBOARD/ANALYTICS ---
@app.route('/api/dashboard', methods=['GET'])
def get_dashboard_stats():
    total_cases = TestCase.query.count()
    status_counts = {}
    for status in TestStatus:
        status_counts[status.value] = TestCase.query.filter_by(status=status.value).count()
    
    priority_counts = {}
    for priority in Priority:
        priority_counts[priority.value] = TestCase.query.filter_by(priority=priority.value).count()
    
    recent_executions = TestCaseExecution.query.order_by(TestCaseExecution.executed_at.desc()).limit(10).all()
    
    return jsonify({
        'total_cases': total_cases,
        'status_counts': status_counts,
        'priority_counts': priority_counts,
        'recent_executions': [
            {
                'id': ex.id,
                'test_case_name': ex.test_case.name,
                'status': ex.status,
                'executed_at': ex.executed_at.isoformat()
            }
            for ex in recent_executions
        ]
    })

# --- API: TEST CASES ---
@app.route('/api/testcases', methods=['GET'])
def get_test_cases():
    # Filtering and searching
    search = request.args.get('search', '')
    status_filter = request.args.get('status', '')
    priority_filter = request.args.get('priority', '')
    category_filter = request.args.get('category', '')
    tag_filter = request.args.get('tag', '')
    
    query = TestCase.query
    
    if search:
        query = query.filter(
            (TestCase.name.contains(search)) | 
            (TestCase.description.contains(search))
        )
    if status_filter:
        query = query.filter_by(status=status_filter)
    if priority_filter:
        query = query.filter_by(priority=priority_filter)
    if category_filter:
        query = query.filter_by(category=category_filter)
    if tag_filter:
        query = query.filter(TestCase.tags.contains(tag_filter))
    
    test_cases = query.order_by(TestCase.created_at.desc()).all()
    
    return jsonify([
        {
            "id": tc.id,
            "name": tc.name,
            "description": tc.description,
            "precondition": tc.precondition or "",
            "postcondition": tc.postcondition or "",
            "comment": tc.comment or "",
            "status": tc.status,
            "priority": tc.priority,
            "category": tc.category or "",
            "tags": tc.tags or "",
            "created_at": tc.created_at.isoformat() if tc.created_at else "",
            "updated_at": tc.updated_at.isoformat() if tc.updated_at else "",
            "steps": [
                {
                    "id": step.id,
                    "description": step.description,
                    "expected_result": step.expected_result,
                    "actual_result": step.actual_result or "",
                    "order": step.order
                }
                for step in sorted(tc.steps, key=lambda s: s.order)
            ],
            "comments_count": len(tc.comments),
            "attachments_count": len(tc.attachments),
            "related_to": tc.related_to
        }
        for tc in test_cases
    ])

@app.route('/api/testcases/<int:test_case_id>', methods=['GET'])
def get_test_case(test_case_id):
    test_case = TestCase.query.get_or_404(test_case_id)
    return jsonify({
        "id": test_case.id,
        "name": test_case.name,
        "description": test_case.description,
        "precondition": test_case.precondition or "",
        "postcondition": test_case.postcondition or "",
        "comment": test_case.comment or "",
        "status": test_case.status,
        "priority": test_case.priority,
        "category": test_case.category or "",
        "tags": test_case.tags or "",
        "created_at": test_case.created_at.isoformat() if test_case.created_at else "",
        "updated_at": test_case.updated_at.isoformat() if test_case.updated_at else "",
        "steps": [
            {
                "id": step.id,
                "description": step.description,
                "expected_result": step.expected_result,
                "actual_result": step.actual_result or "",
                "order": step.order
            }
            for step in sorted(test_case.steps, key=lambda s: s.order)
        ],
        "comments": [
            {
                "id": comment.id,
                "comment": comment.comment,
                "created_at": comment.created_at.isoformat() if comment.created_at else ""
            }
            for comment in test_case.comments
        ],
        "attachments": [
            {
                "id": att.id,
                "filename": att.filename,
                "file_type": att.file_type,
                "created_at": att.created_at.isoformat() if att.created_at else ""
            }
            for att in test_case.attachments
        ],
        "related_to": test_case.related_to,
        "related_cases": [
            {
                "id": rel.id,
                "name": rel.name
            }
            for rel in TestCase.query.filter_by(related_to=test_case.id).all()
        ] if test_case.related_to else []
    })

@app.route('/api/testcases', methods=['POST'])
def create_test_case():
    data = request.json
    
    test_case = TestCase(
        name=data.get('name'),
        description=data.get('description', ''),
        precondition=data.get('precondition', ''),
        postcondition=data.get('postcondition', ''),
        comment=data.get('comment', ''),
        status=data.get('status', TestStatus.NOT_RUN.value),
        priority=data.get('priority', Priority.MEDIUM.value),
        category=data.get('category', ''),
        tags=data.get('tags', ''),
        template_id=data.get('template_id'),
        related_to=data.get('related_to')
    )
    db.session.add(test_case)
    db.session.flush()
    
    # Create version
    version = TestCaseVersion(
        test_case_id=test_case.id,
        version_number=1,
        name=test_case.name,
        description=test_case.description,
        precondition=test_case.precondition,
        postcondition=test_case.postcondition,
        comment=test_case.comment
    )
    db.session.add(version)
    db.session.flush()
    
    # Add steps
    for idx, step_data in enumerate(data.get('steps', [])):
        step = Step(
            test_case_id=test_case.id,
            description=step_data.get('description', ''),
            expected_result=step_data.get('expected_result', ''),
            order=step_data.get('order', idx)
        )
        db.session.add(step)
        
        # Add step to version
        version_step = VersionStep(
            version_id=version.id,
            description=step.description,
            expected_result=step.expected_result,
            order=step.order
        )
        db.session.add(version_step)
    
    db.session.commit()
    return jsonify({"message": "Test Case Created", "id": test_case.id}), 201

@app.route('/api/testcases/<int:test_case_id>', methods=['PUT'])
def update_test_case(test_case_id):
    test_case = TestCase.query.get_or_404(test_case_id)
    data = request.json
    
    # Track changes for versioning
    changed = False
    if data.get('name') != test_case.name or data.get('description') != test_case.description:
        changed = True
    
    test_case.name = data.get('name', test_case.name)
    test_case.description = data.get('description', test_case.description)
    test_case.precondition = data.get('precondition', test_case.precondition)
    test_case.postcondition = data.get('postcondition', test_case.postcondition)
    test_case.comment = data.get('comment', test_case.comment)
    test_case.status = data.get('status', test_case.status)
    test_case.priority = data.get('priority', test_case.priority)
    test_case.category = data.get('category', test_case.category)
    test_case.tags = data.get('tags', test_case.tags)
    test_case.related_to = data.get('related_to', test_case.related_to)
    test_case.updated_at = datetime.utcnow()
    
    # Update steps
    if 'steps' in data:
        # Delete old steps
        Step.query.filter_by(test_case_id=test_case.id).delete()
        
        # Add new steps
        for idx, step_data in enumerate(data['steps']):
            step = Step(
                test_case_id=test_case.id,
                description=step_data.get('description', ''),
                expected_result=step_data.get('expected_result', ''),
                actual_result=step_data.get('actual_result', ''),
                order=step_data.get('order', idx)
            )
            db.session.add(step)
        changed = True
    
    # Create new version if changed
    if changed:
        versions = TestCaseVersion.query.filter_by(test_case_id=test_case.id).order_by(TestCaseVersion.version_number.desc()).all()
        next_version = (versions[0].version_number + 1) if versions else 1
        
        version = TestCaseVersion(
            test_case_id=test_case.id,
            version_number=next_version,
            name=test_case.name,
            description=test_case.description,
            precondition=test_case.precondition,
            postcondition=test_case.postcondition,
            comment=test_case.comment
        )
        db.session.add(version)
        db.session.flush()
        
        # Add steps to version
        for step in test_case.steps:
            version_step = VersionStep(
                version_id=version.id,
                description=step.description,
                expected_result=step.expected_result,
                order=step.order
            )
            db.session.add(version_step)
    
    db.session.commit()
    return jsonify({"message": "Test Case Updated"}), 200

@app.route('/api/testcases/<int:test_case_id>', methods=['DELETE'])
def delete_test_case(test_case_id):
    test_case = TestCase.query.get_or_404(test_case_id)
    db.session.delete(test_case)
    db.session.commit()
    return jsonify({"message": "Test Case Deleted"}), 200

@app.route('/api/testcases/bulk', methods=['POST'])
def bulk_operations():
    data = request.json
    action = data.get('action')
    test_case_ids = data.get('test_case_ids', [])
    
    if action == 'delete':
        TestCase.query.filter(TestCase.id.in_(test_case_ids)).delete(synchronize_session=False)
        db.session.commit()
        return jsonify({"message": f"{len(test_case_ids)} test cases deleted"}), 200
    
    elif action == 'update_status':
        status = data.get('status')
        TestCase.query.filter(TestCase.id.in_(test_case_ids)).update({'status': status}, synchronize_session=False)
        db.session.commit()
        return jsonify({"message": f"Status updated for {len(test_case_ids)} test cases"}), 200
    
    elif action == 'update_priority':
        priority = data.get('priority')
        TestCase.query.filter(TestCase.id.in_(test_case_ids)).update({'priority': priority}, synchronize_session=False)
        db.session.commit()
        return jsonify({"message": f"Priority updated for {len(test_case_ids)} test cases"}), 200
    
    return jsonify({"error": "Invalid action"}), 400

# --- API: COMMENTS ---
@app.route('/api/testcases/<int:test_case_id>/comments', methods=['POST'])
def add_comment(test_case_id):
    test_case = TestCase.query.get_or_404(test_case_id)
    data = request.json
    
    comment = TestCaseComment(
        test_case_id=test_case_id,
        comment=data.get('comment', '')
    )
    db.session.add(comment)
    db.session.commit()
    return jsonify({"message": "Comment added", "id": comment.id}), 201

@app.route('/api/comments/<int:comment_id>', methods=['DELETE'])
def delete_comment(comment_id):
    comment = TestCaseComment.query.get_or_404(comment_id)
    db.session.delete(comment)
    db.session.commit()
    return jsonify({"message": "Comment deleted"}), 200

# --- API: ATTACHMENTS ---
@app.route('/api/testcases/<int:test_case_id>/attachments', methods=['POST'])
def upload_attachment(test_case_id):
    test_case = TestCase.query.get_or_404(test_case_id)
    
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(ATTACHMENT_FOLDER, unique_filename)
        file.save(file_path)
        
        attachment = Attachment(
            test_case_id=test_case_id,
            filename=filename,
            file_path=file_path,
            file_type=filename.rsplit('.', 1)[1].lower()
        )
        db.session.add(attachment)
        db.session.commit()
        return jsonify({"message": "File uploaded", "id": attachment.id}), 201
    
    return jsonify({"error": "Invalid file type"}), 400

@app.route('/api/attachments/<int:attachment_id>', methods=['GET'])
def download_attachment(attachment_id):
    attachment = Attachment.query.get_or_404(attachment_id)
    return send_file(attachment.file_path, as_attachment=True, download_name=attachment.filename)

@app.route('/api/attachments/<int:attachment_id>', methods=['DELETE'])
def delete_attachment(attachment_id):
    attachment = Attachment.query.get_or_404(attachment_id)
    if os.path.exists(attachment.file_path):
        os.remove(attachment.file_path)
    db.session.delete(attachment)
    db.session.commit()
    return jsonify({"message": "Attachment deleted"}), 200

# --- API: TEMPLATES ---
@app.route('/api/templates', methods=['GET'])
def get_templates():
    templates = TestCaseTemplate.query.all()
    return jsonify([
        {
            "id": t.id,
            "name": t.name,
            "description": t.description or "",
            "precondition": t.precondition or "",
            "postcondition": t.postcondition or "",
            "category": t.category or "",
            "steps": [
                {
                    "id": step.id,
                    "description": step.description,
                    "expected_result": step.expected_result,
                    "order": step.order
                }
                for step in sorted(t.template_steps, key=lambda s: s.order)
            ]
        }
        for t in templates
    ])

@app.route('/api/templates', methods=['POST'])
def create_template():
    data = request.json
    
    template = TestCaseTemplate(
        name=data.get('name'),
        description=data.get('description', ''),
        precondition=data.get('precondition', ''),
        postcondition=data.get('postcondition', ''),
        category=data.get('category', '')
    )
    db.session.add(template)
    db.session.flush()
    
    for idx, step_data in enumerate(data.get('steps', [])):
        template_step = TemplateStep(
            template_id=template.id,
            description=step_data.get('description', ''),
            expected_result=step_data.get('expected_result', ''),
            order=step_data.get('order', idx)
        )
        db.session.add(template_step)
    
    db.session.commit()
    return jsonify({"message": "Template created", "id": template.id}), 201

@app.route('/api/templates/<int:template_id>', methods=['DELETE'])
def delete_template(template_id):
    template = TestCaseTemplate.query.get_or_404(template_id)
    db.session.delete(template)
    db.session.commit()
    return jsonify({"message": "Template deleted"}), 200

# --- API: TEST RUNS ---
@app.route('/api/testruns', methods=['GET'])
def get_test_runs():
    test_runs = TestRun.query.order_by(TestRun.created_at.desc()).all()
    return jsonify([
        {
            "id": tr.id,
            "name": tr.name,
            "description": tr.description or "",
            "created_at": tr.created_at.isoformat() if tr.created_at else "",
            "executions_count": len(tr.executions)
        }
        for tr in test_runs
    ])

@app.route('/api/testruns', methods=['POST'])
def create_test_run():
    data = request.json
    
    test_run = TestRun(
        name=data.get('name'),
        description=data.get('description', '')
    )
    db.session.add(test_run)
    db.session.flush()
    
    # Add test cases to run
    for test_case_id in data.get('test_case_ids', []):
        execution = TestCaseExecution(
            test_case_id=test_case_id,
            test_run_id=test_run.id,
            status=TestStatus.NOT_RUN.value
        )
        db.session.add(execution)
    
    db.session.commit()
    return jsonify({"message": "Test run created", "id": test_run.id}), 201

@app.route('/api/testruns/<int:test_run_id>', methods=['GET'])
def get_test_run(test_run_id):
    test_run = TestRun.query.get_or_404(test_run_id)
    return jsonify({
        "id": test_run.id,
        "name": test_run.name,
        "description": test_run.description or "",
        "created_at": test_run.created_at.isoformat() if test_run.created_at else "",
        "executions": [
            {
                "id": ex.id,
                "test_case_id": ex.test_case_id,
                "test_case_name": ex.test_case.name,
                "test_case": {
                    "id": ex.test_case.id,
                    "name": ex.test_case.name,
                    "description": ex.test_case.description,
                    "precondition": ex.test_case.precondition or "",
                    "postcondition": ex.test_case.postcondition or "",
                    "category": ex.test_case.category or "",
                    "priority": ex.test_case.priority,
                    "steps": [
                        {
                            "id": step.id,
                            "description": step.description,
                            "expected_result": step.expected_result,
                            "actual_result": step.actual_result or "",
                            "order": step.order
                        }
                        for step in sorted(ex.test_case.steps, key=lambda s: s.order)
                    ]
                },
                "status": ex.status,
                "executed_at": ex.executed_at.isoformat() if ex.executed_at else "",
                "notes": ex.notes or ""
            }
            for ex in test_run.executions
        ]
    })

@app.route('/api/testruns/<int:test_run_id>', methods=['DELETE'])
def delete_test_run(test_run_id):
    test_run = TestRun.query.get_or_404(test_run_id)
    db.session.delete(test_run)
    db.session.commit()
    return jsonify({"message": "Test run deleted"}), 200

@app.route('/api/testruns/<int:test_run_id>/executions/<int:execution_id>', methods=['PUT'])
def update_execution(test_run_id, execution_id):
    execution = TestCaseExecution.query.get_or_404(execution_id)
    data = request.json
    
    old_status = execution.status
    execution.status = data.get('status', execution.status)
    execution.notes = data.get('notes', execution.notes)
    
    # Update executed_at if status changed from Not Run
    if old_status != execution.status and execution.status != TestStatus.NOT_RUN.value:
        execution.executed_at = datetime.utcnow()
    
    # Update step actual results if provided
    if 'steps' in data:
        for step_data in data['steps']:
            step_id = step_data.get('id')
            if step_id:
                step = Step.query.get(step_id)
                if step:
                    step.actual_result = step_data.get('actual_result', '')
    
    # Update test case status if it's the latest execution
    latest_execution = TestCaseExecution.query.filter_by(test_case_id=execution.test_case_id).order_by(TestCaseExecution.executed_at.desc()).first()
    if latest_execution and latest_execution.id == execution.id:
        execution.test_case.status = execution.status
        execution.test_case.updated_at = datetime.utcnow()
    
    db.session.commit()
    return jsonify({"message": "Execution updated"}), 200

@app.route('/api/testruns/<int:test_run_id>/executions/<int:execution_id>', methods=['DELETE'])
def delete_execution(test_run_id, execution_id):
    execution = TestCaseExecution.query.get_or_404(execution_id)
    db.session.delete(execution)
    db.session.commit()
    return jsonify({"message": "Execution deleted"}), 200

@app.route('/api/steps/<int:step_id>', methods=['PUT'])
def update_step(step_id):
    step = Step.query.get_or_404(step_id)
    data = request.json
    
    step.actual_result = data.get('actual_result', step.actual_result)
    
    db.session.commit()
    return jsonify({"message": "Step updated"}), 200

# --- TEST RUN EXECUTION PAGE ---
@app.route('/testrun/<int:test_run_id>')
def test_run_execution_page(test_run_id):
    return render_template('test_run_execution.html', test_run_id=test_run_id)

# --- API: VERSIONS ---
@app.route('/api/testcases/<int:test_case_id>/versions', methods=['GET'])
def get_versions(test_case_id):
    versions = TestCaseVersion.query.filter_by(test_case_id=test_case_id).order_by(TestCaseVersion.version_number.desc()).all()
    return jsonify([
        {
            "id": v.id,
            "version_number": v.version_number,
            "name": v.name,
            "created_at": v.created_at.isoformat() if v.created_at else "",
            "steps": [
                {
                    "id": step.id,
                    "description": step.description,
                    "expected_result": step.expected_result,
                    "order": step.order
                }
                for step in sorted(v.version_steps, key=lambda s: s.order)
            ]
        }
        for v in versions
    ])

# --- API: EXPORT ---
@app.route('/api/export/<int:test_case_id>', methods=['GET'])
def export_to_word(test_case_id):
    test_case = TestCase.query.get_or_404(test_case_id)
    steps = sorted(test_case.steps, key=lambda s: s.order)
    
    doc = Document()
    doc.add_heading(test_case.name, level=1)
    doc.add_paragraph(f"Description: {test_case.description}")
    
    if test_case.precondition:
        doc.add_paragraph(f"Precondition: {test_case.precondition}")
    if test_case.postcondition:
        doc.add_paragraph(f"Postcondition: {test_case.postcondition}")
    if test_case.comment:
        doc.add_paragraph(f"Comment: {test_case.comment}")
    
    doc.add_paragraph(f"Status: {test_case.status}")
    doc.add_paragraph(f"Priority: {test_case.priority}")
    if test_case.category:
        doc.add_paragraph(f"Category: {test_case.category}")
    if test_case.tags:
        doc.add_paragraph(f"Tags: {test_case.tags}")
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Steps"
    hdr_cells[1].text = "Expected Result"
    hdr_cells[2].text = "Actual Result"
    
    for step in steps:
        row_cells = table.add_row().cells
        row_cells[0].text = step.description
        row_cells[1].text = step.expected_result
        row_cells[2].text = step.actual_result or ""
    
    filename = f"{UPLOAD_FOLDER}/TestCase_{test_case_id}.docx"
    doc.save(filename)
    return send_from_directory(UPLOAD_FOLDER, f"TestCase_{test_case_id}.docx", as_attachment=True)

@app.route('/api/export/bulk', methods=['POST'])
def bulk_export():
    data = request.json
    test_case_ids = data.get('test_case_ids', [])
    
    doc = Document()
    doc.add_heading("Bulk Test Cases Export", level=1)
    
    for test_case_id in test_case_ids:
        test_case = TestCase.query.get(test_case_id)
        if test_case:
            doc.add_heading(test_case.name, level=2)
            doc.add_paragraph(f"Description: {test_case.description}")
            doc.add_paragraph(f"Status: {test_case.status}, Priority: {test_case.priority}")
            
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Steps"
            hdr_cells[1].text = "Expected Result"
            hdr_cells[2].text = "Actual Result"
            
            for step in sorted(test_case.steps, key=lambda s: s.order):
                row_cells = table.add_row().cells
                row_cells[0].text = step.description
                row_cells[1].text = step.expected_result
                row_cells[2].text = step.actual_result or ""
            
            doc.add_page_break()
    
    filename = f"{UPLOAD_FOLDER}/Bulk_Export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return send_from_directory(UPLOAD_FOLDER, os.path.basename(filename), as_attachment=True)

# --- API: IMPORT ---
@app.route('/api/import', methods=['POST'])
def import_test_cases():
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400
    
    try:
        if file.filename.endswith('.xlsx') or file.filename.endswith('.xls'):
            df = pd.read_excel(file)
        elif file.filename.endswith('.csv'):
            df = pd.read_csv(file)
        else:
            return jsonify({"error": "Unsupported file format"}), 400
        
        imported_count = 0
        
        for _, row in df.iterrows():
            test_case = TestCase(
                name=str(row.get('name', '')),
                description=str(row.get('description', '')),
                precondition=str(row.get('precondition', '')),
                postcondition=str(row.get('postcondition', '')),
                status=str(row.get('status', TestStatus.NOT_RUN.value)),
                priority=str(row.get('priority', Priority.MEDIUM.value)),
                category=str(row.get('category', '')),
                tags=str(row.get('tags', ''))
            )
            db.session.add(test_case)
            db.session.flush()
            
            # Import steps if available
            steps_data = row.get('steps', '')
            if steps_data:
                try:
                    steps_list = json.loads(steps_data) if isinstance(steps_data, str) else steps_data
                    for idx, step_data in enumerate(steps_list):
                        step = Step(
                            test_case_id=test_case.id,
                            description=str(step_data.get('description', '')),
                            expected_result=str(step_data.get('expected_result', '')),
                            order=idx
                        )
                        db.session.add(step)
                except:
                    pass
            
            imported_count += 1
        
        db.session.commit()
        return jsonify({"message": f"{imported_count} test cases imported"}), 201
    
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": str(e)}), 400

# --- API: CATEGORIES ---
@app.route('/api/categories', methods=['GET'])
def get_categories():
    categories = db.session.query(TestCase.category).distinct().filter(TestCase.category.isnot(None)).all()
    return jsonify([cat[0] for cat in categories])

# --- API: TAGS ---
@app.route('/api/tags', methods=['GET'])
def get_tags():
    tags_set = set()
    test_cases = TestCase.query.filter(TestCase.tags.isnot(None)).all()
    for tc in test_cases:
        if tc.tags:
            tags_set.update([tag.strip() for tag in tc.tags.split(',')])
    return jsonify(list(tags_set))

# Legacy routes for backward compatibility
@app.route('/testcases', methods=['GET'])
def legacy_get_test_cases():
    return get_test_cases()

@app.route('/testcases', methods=['POST'])
def legacy_create_test_case():
    return create_test_case()

@app.route('/testcases/<int:test_case_id>', methods=['POST'])
def legacy_delete_test_case(test_case_id):
    return delete_test_case(test_case_id)

@app.route('/export/<int:test_case_id>', methods=['GET'])
def legacy_export_to_word(test_case_id):
    return export_to_word(test_case_id)

if __name__ == '__main__':
    app.run(debug=True)
